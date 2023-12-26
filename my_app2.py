from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture(r'C:\Users\ALANE\Downloads\test.png', width=Inches(1.5))

# name phone number and email details
name = input('What is your name? ')
speak('Hello ' + name + ' how are you today?')

speak(name + 'What is your phone number?')
phone_number = input('What is your phone number? ')

email = input('What is your email address? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About Me')
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# More experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company: ')
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)

    else:
        break

# Add skills
skills_detail = input('List your skills: ')
document.add_heading('Skills').italics = True

q = document.add_paragraph(skills_detail).style = 'List Bullet'

# More skills
while True:
    has_more_skills = input('Do you have more skills to list? Yes or No ')
    if has_more_skills == 'yes':
        q = document.add_paragraph()
        q.style = 'List Bullet'
        skills_detail = input('What is the skill? ')
        q.add_run(skills_detail)

    else:
        break

# footer
section = document.sections[0]
footer = section.footer
r = footer.paragraphs[0]
r.text = "I think I did this correctly?????"

document.save('cv.docx')
